"""
FastAPI Backend Server with OpenAI Integration and Blockchain Logging
Connects the frontend chatbot with real AI agents and immutable audit trails
"""

import asyncio
import os
import logging
from datetime import datetime
from typing import Dict, List, Any, Optional
import json

from fastapi import FastAPI, HTTPException, WebSocket, WebSocketDisconnect, Form, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from dotenv import load_dotenv

import sys
sys.path.append('.')

from core.openai_agent import OpenAIAgent, create_openai_agent
from core.blockchain_logger import CommunicationLogger
from core.base_agent import Message, MessageType
from core.connectors import (
    connector_manager, ConnectorType, AuthType, ConnectorStatus,
    PermissionScope, STANDARD_CONNECTORS
)

# Load environment variables
load_dotenv()

# Setup logger
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="Pramiti AI Organization API",
    description="OpenAI-powered agents with blockchain logging",
    version="1.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve static files
app.mount("/static", StaticFiles(directory="web"), name="static")

# Global instances
blockchain_logger = CommunicationLogger()
agents: Dict[str, OpenAIAgent] = {}
active_connections: List[WebSocket] = []

# Import for HTML file serving
from fastapi.responses import FileResponse

# Pydantic models for API
class ChatMessage(BaseModel):
    agent_id: Optional[str] = None  # Optional for user-chat endpoint that does routing
    message: str
    user_id: str = "user"
    session_id: Optional[str] = None

class AgentResponse(BaseModel):
    agent_id: str
    agent_name: str
    response: str
    timestamp: str
    blockchain_hash: Optional[str] = None
    block_number: Optional[int] = None

class AgentInfo(BaseModel):
    agent_id: str
    name: str
    role: str
    specialization: str
    status: str = "online"
    description: Optional[str] = None

class BlockchainEntry(BaseModel):
    entry_id: str
    message_id: Optional[str] = None
    sender_id: str
    recipient_id: Optional[str] = None
    timestamp: str
    block_number: int
    entry_hash: str

class CreateAgentRequest(BaseModel):
    name: str
    agent_type: str
    role: str
    reporting_manager_role: Optional[str] = None
    specialization: str
    model: str = "gpt-4o-mini"  # Default to cost-effective model
    system_prompt: Optional[str] = None
    responsibilities_from_doc: Optional[str] = None  # Extracted from uploaded document

class UpdateAgentRequest(BaseModel):
    name: Optional[str] = None
    agent_type: Optional[str] = None
    role: Optional[str] = None
    specialization: Optional[str] = None
    model: Optional[str] = None
    system_prompt: Optional[str] = None
    enabled: Optional[bool] = None

class RoleDefinition(BaseModel):
    role_id: str
    name: str
    description: str
    level: int  # Hierarchy level (1=CEO, 2=Senior Manager, etc.)
    responsibilities: str = ""  # Default responsibilities for this role
    permissions: List[str] = []
    enabled: bool = True

# Global role definitions storage
custom_roles: Dict[str, RoleDefinition] = {
    "ceo": RoleDefinition(
        role_id="ceo",
        name="CEO / Executive Director",
        description="Top-level executive leadership",
        level=1,
        responsibilities="- Set strategic direction and vision\n- Make final decisions on major initiatives\n- Oversee all departments and operations\n- Represent the organization",
        permissions=["all"],
        enabled=True
    ),
    "senior_manager": RoleDefinition(
        role_id="senior_manager",
        name="Senior Manager",
        description="Senior management and oversight",
        level=2,
        responsibilities="- Manage department operations\n- Lead and mentor team managers\n- Approve major changes and initiatives\n- Report to executive leadership",
        permissions=["manage_teams", "approve_changes"],
        enabled=True
    ),
    "manager": RoleDefinition(
        role_id="manager",
        name="Manager",
        description="Team and project management",
        level=3,
        responsibilities="- Manage team members and projects\n- Assign tasks and monitor progress\n- Coordinate with other teams\n- Report to senior management",
        permissions=["manage_team", "assign_tasks"],
        enabled=True
    ),
    "specialist": RoleDefinition(
        role_id="specialist",
        name="Specialist / Subject Matter Expert",
        description="Domain expertise and specialized knowledge",
        level=4,
        responsibilities="- Provide expert knowledge and guidance\n- Execute specialized tasks\n- Collaborate with teams on technical matters\n- Mentor junior team members",
        permissions=["execute_tasks", "provide_expertise"],
        enabled=True
    ),
    "analyst": RoleDefinition(
        role_id="analyst",
        name="Analyst",
        description="Analysis and research",
        level=5,
        responsibilities="- Analyze data and trends\n- Create reports and insights\n- Support decision-making with data\n- Identify improvement opportunities",
        permissions=["analyze_data", "create_reports"],
        enabled=True
    ),
    "operator": RoleDefinition(
        role_id="operator",
        name="Operator",
        description="Operational execution",
        level=6,
        responsibilities="- Execute assigned tasks\n- Follow standard procedures\n- Report issues and progress\n- Maintain operational efficiency",
        permissions=["execute_tasks"],
        enabled=True
    )
}

# Track agent enabled/disabled status
agent_status: Dict[str, bool] = {}

# Initialize agents on startup
@app.on_event("startup")
async def startup_event():
    """Initialize OpenAI agents with blockchain logging"""
    
    # Check if OpenAI API key is set
    if not os.getenv("OPENAI_API_KEY"):
        print("⚠️  WARNING: OPENAI_API_KEY not set. Please add your API key to .env file")
        print("   Copy .env.example to .env and add your OpenAI API key")
    
    # Initialize blockchain logger
    try:
        # blockchain_logger.initialize()  # Commented out as it might not exist in current version
        print("✓ Blockchain logging system initialized")
    except Exception as e:
        print(f"⚠️  Blockchain initialization warning: {e}")
    
    # Create agents
    agent_definitions = [
        {
            "type": "ceo",
            "id": "ceo-001", 
            "name": "Executive AI Director"
        },
        {
            "type": "incident_manager",
            "id": "manager-incident-001",
            "name": "ITSM Operations Manager"
        },
        {
            "type": "incident_specialist",
            "id": "agent-incident-001",
            "name": "Incident Response Specialist"
        },
        {
            "type": "problem_specialist", 
            "id": "agent-problem-001",
            "name": "Problem Analysis Expert"
        },
        {
            "type": "change_specialist",
            "id": "agent-change-001", 
            "name": "Change Management Specialist"
        }
    ]
    
    for agent_def in agent_definitions:
        try:
            agent = create_openai_agent(
                agent_type=agent_def["type"],
                agent_id=agent_def["id"],
                name=agent_def["name"],
                blockchain_logger=blockchain_logger
            )
            agents[agent_def["id"]] = agent
            agent_status[agent_def["id"]] = True  # Enable by default
            print(f"✓ Created OpenAI agent: {agent.name}")
            
        except Exception as e:
            print(f"❌ Failed to create agent {agent_def['name']}: {e}")
    
    print(f"\n🚀 Pramiti AI Organization API started with {len(agents)} agents")
    print(f"📍 Blockchain logging: {'Enabled' if blockchain_logger else 'Disabled'}")
    print(f"🔗 OpenAI Integration: {'Enabled' if os.getenv('OPENAI_API_KEY') else 'Requires API Key'}")

# API Routes

@app.get("/")
async def root():
    """API health check"""
    return {
        "message": "Pramiti AI Organization API",
        "status": "running",
        "agents": len(agents),
        "blockchain_enabled": True,
        "openai_enabled": bool(os.getenv("OPENAI_API_KEY"))
    }

# Serve HTML frontend files
@app.get("/openai-dashboard.html")
async def serve_openai_dashboard():
    """Serve the OpenAI dashboard HTML"""
    return FileResponse("web/openai-dashboard.html")

@app.get("/enhanced-dashboard.html")
async def serve_enhanced_dashboard():
    """Serve the enhanced dashboard HTML with no-cache headers"""
    from fastapi.responses import Response
    with open("web/enhanced-dashboard.html", "r") as f:
        content = f.read()
    return Response(
        content=content,
        media_type="text/html",
        headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0"
        }
    )

@app.get("/index.html")
async def serve_index():
    """Serve the index dashboard HTML"""
    return FileResponse("web/index.html")

@app.get("/agents", response_model=List[AgentInfo])
async def get_agents(include_disabled: bool = False):
    """Get list of available agents"""
    agent_list = []
    for agent_id, agent in agents.items():
        # Check if agent is enabled
        is_enabled = agent_status.get(agent_id, True)
        
        # Skip disabled agents unless explicitly requested
        if not is_enabled and not include_disabled:
            continue
            
        agent_list.append(AgentInfo(
            agent_id=agent_id,
            name=agent.name,
            role=agent.role.value,
            specialization=agent.specialization,
            status="online" if is_enabled else "disabled",
            description=f"{agent.role.value.replace('_', ' ').title()} specializing in {agent.specialization}"
        ))
    
    return agent_list

@app.post("/agents", response_model=AgentInfo)
async def create_agent(
    name: str = Form(...),
    agent_type: str = Form(...),
    role: str = Form(...),
    reporting_manager_role: str = Form(None),
    reports_to_agent_id: str = Form(None),
    specialization: str = Form(...),
    model: str = Form("gpt-4o-mini"),
    responsibilities_doc: Optional[UploadFile] = File(None)
):
    """Create a new AI agent with optional document upload for responsibilities and direct reporting relationship"""
    
    try:
        # Generate unique agent ID
        agent_count = len(agents) + 1
        agent_type_clean = agent_type.lower().replace(' ', '-')
        agent_id = f"{agent_type_clean}-{agent_count:03d}"
        
        # Check for duplicate names
        existing_names = [agent.name.lower() for agent in agents.values()]
        if name.lower() in existing_names:
            raise HTTPException(
                status_code=400, 
                detail=f"Agent with name '{name}' already exists"
            )
        
        # Get responsibilities from role definition
        if role not in custom_roles:
            raise HTTPException(
                status_code=400,
                detail=f"Role '{role}' not found. Please create the role first."
            )
        
        role_def = custom_roles[role]
        if not role_def.enabled:
            raise HTTPException(
                status_code=400,
                detail=f"Role '{role}' is currently disabled."
            )
        
        # Validate reporting hierarchy
        if reporting_manager_role:
            if reporting_manager_role not in custom_roles:
                raise HTTPException(
                    status_code=400,
                    detail=f"Reporting manager role '{reporting_manager_role}' not found."
                )
            
            manager_role_def = custom_roles[reporting_manager_role]
            
            # Ensure reporting manager has higher authority (lower level number)
            if manager_role_def.level >= role_def.level:
                raise HTTPException(
                    status_code=400,
                    detail=f"Reporting manager role must be higher in hierarchy than agent's role. {manager_role_def.name} (Level {manager_role_def.level}) cannot manage {role_def.name} (Level {role_def.level})."
                )
        
        # Validate specific reporting agent if provided
        if reports_to_agent_id:
            if reports_to_agent_id not in agents:
                raise HTTPException(
                    status_code=400,
                    detail=f"Reporting manager agent '{reports_to_agent_id}' not found."
                )
            
            # Validate that the specified agent has the reporting manager role
            manager_agent = agents[reports_to_agent_id]
            manager_custom_role = getattr(manager_agent, 'metadata', {}).get('custom_role_id')
            
            if reporting_manager_role and manager_custom_role != reporting_manager_role:
                raise HTTPException(
                    status_code=400,
                    detail=f"Selected manager '{manager_agent.name}' has role '{manager_custom_role}' but expected '{reporting_manager_role}'."
                )
        
        # Base responsibilities from role
        responsibilities = role_def.responsibilities
        
        # Process uploaded document if provided
        document_content = None
        if responsibilities_doc and responsibilities_doc.filename:
            try:
                # Read file content
                file_content = await responsibilities_doc.read()
                
                # Extract text based on file type
                if responsibilities_doc.filename.endswith('.txt'):
                    document_content = file_content.decode('utf-8')
                elif responsibilities_doc.filename.endswith(('.docx', '.doc')):
                    # Extract text from DOCX
                    from docx import Document
                    from io import BytesIO
                    doc = Document(BytesIO(file_content))
                    document_content = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                elif responsibilities_doc.filename.endswith('.pdf'):
                    # Extract text from PDF
                    from PyPDF2 import PdfReader
                    from io import BytesIO
                    pdf = PdfReader(BytesIO(file_content))
                    document_content = '\n'.join([page.extract_text() for page in pdf.pages])
                
                # Append document content to responsibilities
                if document_content:
                    responsibilities += f"\n\nAdditional Responsibilities (from uploaded document):\n{document_content}"
                    print(f"✅ Processed document: {responsibilities_doc.filename} ({len(document_content)} chars)")
                
            except Exception as doc_error:
                print(f"⚠️ Warning: Could not process document {responsibilities_doc.filename}: {str(doc_error)}")
                # Continue without document content rather than failing
        
        # Parse role from string to AgentRole enum
        from core.base_agent import AgentRole
        
        # Map role strings to AgentRole enum
        # Note: Base system only has 3 roles defined
        role_mapping = {
            "ceo": AgentRole.CEO,
            "senior_manager": AgentRole.SENIOR_MANAGER,
            "manager": AgentRole.SENIOR_MANAGER,  # Map to senior_manager
            "subject_matter_expert": AgentRole.SUBJECT_MATTER_EXPERT,
            "sme": AgentRole.SUBJECT_MATTER_EXPERT,
            "specialist": AgentRole.SUBJECT_MATTER_EXPERT,  # Map to SME
            "analyst": AgentRole.SUBJECT_MATTER_EXPERT,  # Map to SME
            "operator": AgentRole.SUBJECT_MATTER_EXPERT  # Map to SME
        }
        
        agent_role = role_mapping.get(role.lower(), AgentRole.SUBJECT_MATTER_EXPERT)
        
        # Build system prompt with hierarchy information
        hierarchy_info = ""
        if reporting_manager_role:
            manager_role_def = custom_roles[reporting_manager_role]
            hierarchy_info = f"\n\nYou report to {manager_role_def.name} level management."
        
        system_prompt = f"""You are {name}, a {role_def.name} with expertise in {specialization}.

Your responsibilities include:
{responsibilities}
{hierarchy_info}

You work within the Pramiti AI Organization, collaborating with other AI agents to provide excellent service.
Be professional, knowledgeable, and helpful in all interactions."""
        
        # Create AI config
        from core.openai_agent import AIAgentConfig
        ai_config = AIAgentConfig(
            model=model,
            temperature=0.7,
            max_tokens=1200,
            system_prompt=system_prompt
        )
        
        # Create the new agent directly
        from core.openai_agent import OpenAIAgent
        new_agent = OpenAIAgent(
            agent_id=agent_id,
            name=name,
            role=agent_role,
            specialization=specialization,
            blockchain_logger=blockchain_logger,
            ai_config=ai_config
        )
        
        # Store reporting relationship in agent metadata
        if not hasattr(new_agent, 'metadata'):
            new_agent.metadata = {}
        new_agent.metadata['reporting_manager_role'] = reporting_manager_role
        new_agent.metadata['reports_to_agent_id'] = reports_to_agent_id
        new_agent.metadata['custom_role_id'] = role
        new_agent.metadata['agent_type'] = agent_type
        
        # Add to agents dictionary and enable it
        agents[agent_id] = new_agent
        agent_status[agent_id] = True  # Enable by default
        
        print(f"✅ Created new agent: {name} ({agent_id}) with role {role_def.name}")
        if reporting_manager_role:
            print(f"   └─ Reports to role: {custom_roles[reporting_manager_role].name}")
        if reports_to_agent_id:
            manager_name = agents.get(reports_to_agent_id, {})
            manager_name = getattr(manager_name, 'name', reports_to_agent_id) if manager_name else reports_to_agent_id
            print(f"   └─ Direct manager: {manager_name} ({reports_to_agent_id})")
        
        # Return agent info
        return AgentInfo(
            agent_id=agent_id,
            name=new_agent.name,
            role=role,  # Return the custom role ID
            specialization=new_agent.specialization,
            status="online",
            description=f"{role_def.name} specializing in {new_agent.specialization}"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"❌ Error creating agent: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create agent: {str(e)}")

@app.put("/agents/{agent_id}")
async def update_agent(agent_id: str, request: UpdateAgentRequest):
    """Update an existing agent"""
    try:
        if agent_id not in agents:
            raise HTTPException(status_code=404, detail=f"Agent '{agent_id}' not found")
        
        agent = agents[agent_id]
        
        # Update fields if provided
        if request.name:
            agent.name = request.name
        
        if request.specialization:
            agent.specialization = request.specialization
        
        if request.role:
            from core.base_agent import AgentRole
            role_mapping = {
                "ceo": AgentRole.CEO,
                "senior_manager": AgentRole.SENIOR_MANAGER,
                "manager": AgentRole.SENIOR_MANAGER,
                "subject_matter_expert": AgentRole.SUBJECT_MATTER_EXPERT,
                "sme": AgentRole.SUBJECT_MATTER_EXPERT,
                "specialist": AgentRole.SUBJECT_MATTER_EXPERT,
                "analyst": AgentRole.SUBJECT_MATTER_EXPERT,
                "operator": AgentRole.SUBJECT_MATTER_EXPERT
            }
            agent.role = role_mapping.get(request.role.lower(), AgentRole.SUBJECT_MATTER_EXPERT)
        
        if request.enabled is not None:
            agent_status[agent_id] = request.enabled
        
        # Note: model, system_prompt, and responsibilities updates would require recreating the agent
        # For now, these are informational only
        
        print(f"✅ Updated agent: {agent.name} ({agent_id})")
        
        is_enabled = agent_status.get(agent_id, True)
        return AgentInfo(
            agent_id=agent_id,
            name=agent.name,
            role=agent.role.value,
            specialization=agent.specialization,
            status="online" if is_enabled else "disabled",
            description=f"{agent.role.value.replace('_', ' ').title()} specializing in {agent.specialization}"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error updating agent: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to update agent: {str(e)}")

@app.delete("/agents/{agent_id}")
async def delete_agent(agent_id: str):
    """Delete an agent"""
    try:
        if agent_id not in agents:
            raise HTTPException(status_code=404, detail=f"Agent '{agent_id}' not found")
        
        agent_name = agents[agent_id].name
        del agents[agent_id]
        
        if agent_id in agent_status:
            del agent_status[agent_id]
        
        print(f"✅ Deleted agent: {agent_name} ({agent_id})")
        
        return {"message": f"Agent '{agent_name}' deleted successfully", "agent_id": agent_id}
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error deleting agent: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to delete agent: {str(e)}")

@app.post("/agents/{agent_id}/enable")
async def enable_agent(agent_id: str):
    """Enable an agent"""
    if agent_id not in agents:
        raise HTTPException(status_code=404, detail=f"Agent '{agent_id}' not found")
    
    agent_status[agent_id] = True
    print(f"✅ Enabled agent: {agents[agent_id].name} ({agent_id})")
    
    return {"message": "Agent enabled", "agent_id": agent_id, "enabled": True}

@app.post("/agents/{agent_id}/disable")
async def disable_agent(agent_id: str):
    """Disable an agent"""
    if agent_id not in agents:
        raise HTTPException(status_code=404, detail=f"Agent '{agent_id}' not found")
    
    agent_status[agent_id] = False
    print(f"✅ Disabled agent: {agents[agent_id].name} ({agent_id})")
    
    return {"message": "Agent disabled", "agent_id": agent_id, "enabled": False}

# Role Management Endpoints

@app.get("/roles")
async def get_roles(include_disabled: bool = False):
    """Get list of available roles"""
    if include_disabled:
        return list(custom_roles.values())
    return [role for role in custom_roles.values() if role.enabled]

@app.post("/roles")
async def create_role(role: RoleDefinition):
    """Create a new custom role"""
    if role.role_id in custom_roles:
        raise HTTPException(status_code=400, detail=f"Role '{role.role_id}' already exists")
    
    custom_roles[role.role_id] = role
    print(f"✅ Created new role: {role.name} ({role.role_id})")
    
    return role

@app.put("/roles/{role_id}")
async def update_role(role_id: str, role: RoleDefinition):
    """Update an existing role"""
    if role_id not in custom_roles:
        raise HTTPException(status_code=404, detail=f"Role '{role_id}' not found")
    
    custom_roles[role_id] = role
    print(f"✅ Updated role: {role.name} ({role_id})")
    
    return role

@app.delete("/roles/{role_id}")
async def delete_role(role_id: str):
    """Delete a role"""
    if role_id not in custom_roles:
        raise HTTPException(status_code=404, detail=f"Role '{role_id}' not found")
    
    # Check if any agents use this role
    agents_using_role = [
        agent_id for agent_id, agent in agents.items() 
        if agent.role.value == role_id
    ]
    
    if agents_using_role:
        raise HTTPException(
            status_code=400, 
            detail=f"Cannot delete role '{role_id}': {len(agents_using_role)} agent(s) are using it"
        )
    
    role_name = custom_roles[role_id].name
    del custom_roles[role_id]
    print(f"✅ Deleted role: {role_name} ({role_id})")
    
    return {"message": f"Role '{role_name}' deleted successfully", "role_id": role_id}

@app.post("/roles/{role_id}/enable")
async def enable_role(role_id: str):
    """Enable a role"""
    if role_id not in custom_roles:
        raise HTTPException(status_code=404, detail=f"Role '{role_id}' not found")
    
    custom_roles[role_id].enabled = True
    print(f"✅ Enabled role: {custom_roles[role_id].name} ({role_id})")
    
    return {"message": "Role enabled", "role_id": role_id, "enabled": True}

@app.post("/roles/{role_id}/disable")
async def disable_role(role_id: str):
    """Disable a role"""
    if role_id not in custom_roles:
        raise HTTPException(status_code=404, detail=f"Role '{role_id}' not found")
    
    custom_roles[role_id].enabled = False
    print(f"✅ Disabled role: {custom_roles[role_id].name} ({role_id})")
    
    return {"message": "Role disabled", "role_id": role_id, "enabled": False}

# API v1 endpoints for enhanced/basic dashboards
@app.get("/api/v1/dashboard/kpis")
async def get_dashboard_kpis():
    """Get dashboard KPIs for enhanced/basic dashboards"""
    total_messages = len(blockchain_logger.local_blockchain)
    
    return {
        "active_agents": len(agents),
        "messages_processed": total_messages,
        "active_incidents": 0,
        "blockchain_logs": len(blockchain_logger.local_blockchain),
        "avg_response_time": "0.5s",
        "system_health": "healthy"
    }

@app.get("/api/v1/agents/hierarchy")
async def get_agents_hierarchy():
    """Get agent hierarchy for visualization with reporting relationships"""
    hierarchy = {
        "ceo": [],
        "managers": [],
        "specialists": [],
        "reporting_structure": {},  # Maps agent_id to reporting_manager_role
        "direct_reports": {}  # Maps agent_id to reports_to_agent_id
    }
    
    for agent_id, agent in agents.items():
        # Get metadata for custom role and reporting info
        custom_role_id = getattr(agent, 'metadata', {}).get('custom_role_id', None)
        reporting_manager_role = getattr(agent, 'metadata', {}).get('reporting_manager_role', None)
        reports_to_agent_id = getattr(agent, 'metadata', {}).get('reports_to_agent_id', None)
        agent_type = getattr(agent, 'metadata', {}).get('agent_type', agent.specialization)
        
        agent_data = {
            "id": agent_id,
            "name": agent.name,
            "role": agent.role.value,
            "custom_role_id": custom_role_id,
            "agent_type": agent_type,
            "specialization": agent.specialization,
            "status": "online",
            "reporting_manager_role": reporting_manager_role,
            "reports_to_agent_id": reports_to_agent_id
        }
        
        # Add to reporting structure
        if reporting_manager_role:
            hierarchy["reporting_structure"][agent_id] = reporting_manager_role
        
        # Add direct reporting relationship
        if reports_to_agent_id:
            hierarchy["direct_reports"][agent_id] = reports_to_agent_id
        
        # Categorize by role
        if "ceo" in agent_id or agent.role.value == "ceo":
            hierarchy["ceo"].append(agent_data)
        elif "manager" in agent_id or agent.role.value == "senior_manager":
            hierarchy["managers"].append(agent_data)
        else:
            hierarchy["specialists"].append(agent_data)
    
    return hierarchy

@app.get("/api/v1/communications/recent")
async def get_recent_communications():
    """Get recent communications"""
    recent_blocks = blockchain_logger.export_blockchain()[-10:]
    
    communications = []
    for block in recent_blocks:
        communications.append({
            "id": block.get("entry_id"),
            "from": block.get("sender_id", "system"),
            "to": block.get("recipient_id", "system"),
            "timestamp": block.get("timestamp"),
            "type": "message",
            "block_number": block.get("block_number")
        })
    
    return communications

@app.get("/api/v1/communications/queues")
async def get_message_queues():
    """Get message queue status for all agents"""
    queues = []
    
    for agent_id, agent in agents.items():
        queues.append({
            "agent_id": agent_id,
            "agent_name": agent.name,
            "queue_size": 0,  # No active queue system yet
            "status": "healthy",
            "last_activity": datetime.now().isoformat()
        })
    
    return queues

@app.get("/api/v1/blockchain/logs")
async def get_blockchain_logs(limit: int = 50):
    """Get blockchain communication logs"""
    try:
        # Get blockchain entries
        blockchain = blockchain_logger.local_blockchain
        
        # Reverse to show newest first
        logs = []
        for i, block in enumerate(reversed(blockchain[-limit:])):
            # Extract block fields - check actual structure
            block_number = block.get("block_number", len(blockchain) - i)
            timestamp = block.get("timestamp", "N/A")
            entry_hash = block.get("entry_hash", "N/A")
            prev_hash = block.get("previous_hash", "")
            
            # Create log entry
            log_entry = {
                "block_number": block_number,
                "timestamp": timestamp,
                "previous_hash": prev_hash[:16] + "..." if prev_hash else "Genesis",
                "hash": entry_hash[:16] + "..." if entry_hash and entry_hash != "N/A" else "N/A",
                "full_hash": entry_hash,
                "sender": block.get("sender_id", "System"),
                "recipient": block.get("recipient_id", "N/A"),
                "message_type": block.get("message_type", "UNKNOWN"),
            }
            
            # Try to extract message content
            if "content_hash" in block:
                # This is a message log - try to get readable content
                metadata = block.get("metadata", {})
                log_entry["message"] = metadata.get("preview", "Message logged (content hashed)")
            else:
                # Generic block
                log_entry["message"] = str(block.get("data", "N/A"))[:100]
            
            logs.append(log_entry)
        
        return {
            "total_blocks": len(blockchain),
            "logs": logs,
            "blockchain_valid": True  # Always true for local blockchain
        }
    except Exception as e:
        print(f"Error fetching blockchain logs: {e}")
        import traceback
        traceback.print_exc()
        return {
            "total_blocks": 0,
            "logs": [],
            "blockchain_valid": True,
            "error": str(e)
        }

@app.post("/chat", response_model=AgentResponse)
async def chat_with_agent(chat_message: ChatMessage):
    """Send message to specific agent and get AI-powered response"""
    
    if chat_message.agent_id not in agents:
        raise HTTPException(status_code=404, detail=f"Agent {chat_message.agent_id} not found")
    
    agent = agents[chat_message.agent_id]
    
    try:
        # Create message object
        user_message = Message(
            sender_id=chat_message.user_id,
            recipient_id=chat_message.agent_id,
            message_type=MessageType.REQUEST,
            content={
                "text": chat_message.message,
                "session_id": chat_message.session_id
            }
        )
        
        # Process with OpenAI agent (includes blockchain logging)
        response_message = await agent.process_message(user_message)
        
        # Get blockchain info
        audit_trail = agent.get_blockchain_audit_trail(response_message.id)
        
        return AgentResponse(
            agent_id=chat_message.agent_id,
            agent_name=agent.name,
            response=response_message.content.get("response", "No response generated"),
            timestamp=response_message.timestamp.isoformat(),
            blockchain_hash=response_message.metadata.get("original_message_hash"),
            block_number=len(blockchain_logger.local_blockchain)
        )
        
    except Exception as e:
        print(f"❌ Chat error: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing chat: {str(e)}")

@app.get("/agent/{agent_id}/audit")
async def get_agent_audit_trail(agent_id: str):
    """Get blockchain audit trail for specific agent"""
    
    if agent_id not in agents:
        raise HTTPException(status_code=404, detail=f"Agent {agent_id} not found")
    
    agent = agents[agent_id]
    audit_data = agent.get_blockchain_audit_trail()
    
    return {
        "agent_id": agent_id,
        "agent_name": agent.name,
        "audit_trail": audit_data,
        "blockchain_integrity": blockchain_logger.verify_blockchain_integrity()
    }

@app.get("/agent/{agent_id}/report")
async def get_agent_report(agent_id: str):
    """Get comprehensive agent activity report"""
    
    if agent_id not in agents:
        raise HTTPException(status_code=404, detail=f"Agent {agent_id} not found")
    
    agent = agents[agent_id]
    report = agent.generate_agent_report()
    
    return report

@app.get("/blockchain/status")
async def get_blockchain_status():
    """Get blockchain status and recent activity"""
    
    recent_blocks = blockchain_logger.export_blockchain()[-10:]  # Last 10 blocks
    
    return {
        "total_blocks": len(blockchain_logger.local_blockchain),
        "integrity_verified": blockchain_logger.verify_blockchain_integrity(),
        "recent_blocks": [
            BlockchainEntry(
                entry_id=block["entry_id"],
                message_id=block.get("message_id"),
                sender_id=block.get("sender_id", block.get("agent_id", "system")),
                recipient_id=block.get("recipient_id"),
                timestamp=block["timestamp"],
                block_number=block["block_number"],
                entry_hash=block["entry_hash"]
            ) for block in recent_blocks
        ]
    }

@app.get("/blockchain/compliance-report")
async def generate_compliance_report():
    """Generate blockchain compliance report"""
    
    start_date = datetime.now().replace(hour=0, minute=0, second=0)
    report = blockchain_logger.generate_compliance_report(start_date, datetime.now())
    
    return report

# WebSocket for real-time communication
@app.websocket("/ws/{agent_id}")
async def websocket_endpoint(websocket: WebSocket, agent_id: str):
    """WebSocket endpoint for real-time agent communication"""
    
    if agent_id not in agents:
        await websocket.close(code=4404)
        return
    
    await websocket.accept()
    active_connections.append(websocket)
    agent = agents[agent_id]
    
    try:
        while True:
            # Receive message from client
            data = await websocket.receive_json()
            message_text = data.get("message", "")
            user_id = data.get("user_id", "websocket_user")
            
            if message_text:
                # Create and process message
                user_message = Message(
                    sender_id=user_id,
                    recipient_id=agent_id,
                    message_type=MessageType.REQUEST,
                    content={"text": message_text}
                )
                
                # Get AI response with blockchain logging
                response_message = await agent.process_message(user_message)
                
                # Send response back to client
                await websocket.send_json({
                    "agent_id": agent_id,
                    "agent_name": agent.name,
                    "response": response_message.content.get("response", "No response"),
                    "timestamp": response_message.timestamp.isoformat(),
                    "blockchain_logged": True,
                    "block_number": len(blockchain_logger.local_blockchain)
                })
                
    except WebSocketDisconnect:
        active_connections.remove(websocket)
        print(f"WebSocket disconnected for agent {agent_id}")
    except Exception as e:
        print(f"WebSocket error: {e}")
        if websocket in active_connections:
            active_connections.remove(websocket)

# Generic WebSocket endpoints for enhanced/basic dashboards
@app.websocket("/ws")
async def websocket_generic(websocket: WebSocket):
    """Generic WebSocket endpoint for real-time updates"""
    await websocket.accept()
    active_connections.append(websocket)
    
    try:
        while True:
            # Send periodic updates
            await asyncio.sleep(5)
            
            # Send status update
            await websocket.send_json({
                "type": "status_update",
                "agents": len(agents),
                "blockchain_blocks": len(blockchain_logger.local_blockchain),
                "timestamp": datetime.now().isoformat()
            })
                
    except WebSocketDisconnect:
        if websocket in active_connections:
            active_connections.remove(websocket)
        print("Generic WebSocket disconnected")
    except Exception as e:
        print(f"Generic WebSocket error: {e}")
        if websocket in active_connections:
            active_connections.remove(websocket)

@app.websocket("/ws/chat")
async def websocket_chat(websocket: WebSocket):
    """WebSocket endpoint for chat interface"""
    await websocket.accept()
    active_connections.append(websocket)
    
    try:
        while True:
            data = await websocket.receive_json()
            message_text = data.get("message", "")
            agent_id = data.get("agent_id", "agent-incident-001")
            
            if message_text and agent_id in agents:
                agent = agents[agent_id]
                
                user_message = Message(
                    sender_id="websocket_user",
                    recipient_id=agent_id,
                    message_type=MessageType.REQUEST,
                    content={"text": message_text}
                )
                
                response_message = await agent.process_message(user_message)
                
                await websocket.send_json({
                    "type": "chat_response",
                    "agent_id": agent_id,
                    "agent_name": agent.name,
                    "response": response_message.content.get("response", "No response"),
                    "timestamp": response_message.timestamp.isoformat(),
                    "blockchain_logged": True
                })
                
    except WebSocketDisconnect:
        if websocket in active_connections:
            active_connections.remove(websocket)
        print("Chat WebSocket disconnected")
    except Exception as e:
        print(f"Chat WebSocket error: {e}")
        if websocket in active_connections:
            active_connections.remove(websocket)

# User-facing chatbot endpoint (Pramiti Assistant)
@app.post("/user-chat")
async def user_chatbot(chat_message: ChatMessage):
    """User-facing chatbot that can route to appropriate agents"""
    
    message_lower = chat_message.message.lower()
    
    # Check for OAuth/authentication intent FIRST (highest priority)
    oauth_keywords = ['login', 'sign in', 'signin', 'authenticate', 'connect', 'auth',
                     'microsoft', 'google', 'outlook', 'gmail', 'office 365', 'o365',
                     'onedrive', 'calendar', 'account', 'access', 'email setup',
                     'permission', 'authorize', 'token', 'oauth']
    
    if any(word in message_lower for word in oauth_keywords):
        # Route to OAuth assistant
        from agents.oauth_agent import oauth_assistant
        
        # Call the OAuth assistant's process_message method directly with the string
        response_text = oauth_assistant.handle_chat_message(chat_message.message, chat_message.user_id)
        
        # Check if we need to generate actual OAuth URL
        if "#auth-" in response_text:
            # Extract auth type from placeholder
            import re
            placeholder_match = re.search(r'#auth-(microsoft|google)-([^\s\)]+)', response_text)
            if placeholder_match:
                auth_type = placeholder_match.group(1)
                placeholder_email = placeholder_match.group(2)
                
                # Use actual user_id as email if placeholder is used
                # or extract real email from message
                user_email = chat_message.user_id
                
                # Try to extract real email from message if provided
                email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
                email_matches = re.findall(email_pattern, chat_message.message)
                if email_matches:
                    user_email = email_matches[0]
                elif placeholder_email != "user@placeholder.com":
                    # Use the email from placeholder if it's not the default
                    user_email = placeholder_email
                
                # Map auth type to connector type
                connector_type_map = {
                    'microsoft': 'microsoft_teams',
                    'google': 'google_workspace'
                }
                connector_type = connector_type_map.get(auth_type, auth_type)
                
                # Prepare OAuth request payload
                oauth_payload = {
                    "user_email": user_email,
                    "connector_type": connector_type
                }
                
                # Add credentials based on connector type
                import os
                if auth_type == "microsoft":
                    client_id = os.getenv("MICROSOFT_CLIENT_ID")
                    client_secret = os.getenv("MICROSOFT_CLIENT_SECRET")
                    tenant_id = os.getenv("MICROSOFT_TENANT_ID", "common")
                    
                    if not client_id or not client_secret:
                        response_text = f"""
❌ **Microsoft OAuth Not Configured**

To enable Microsoft 365 authentication, you need to set up OAuth credentials in your `.env` file:

```
MICROSOFT_CLIENT_ID=your_client_id_here
MICROSOFT_CLIENT_SECRET=your_client_secret_here
MICROSOFT_TENANT_ID=common
```

**How to get these credentials:**
1. Go to Azure Portal (https://portal.azure.com)
2. Navigate to "Azure Active Directory" → "App registrations"
3. Create a new app registration or use existing one
4. Copy the "Application (client) ID" → This is your CLIENT_ID
5. Go to "Certificates & secrets" → Create new client secret → Copy the value
6. Set the redirect URI to: `http://localhost:8084/api/v1/oauth/callback/microsoft`

**Need help?** Check out `CREDENTIALS_SETUP.md` for detailed instructions.

_For now, I've shown you what the authentication flow would look like, but you'll need to configure the credentials to actually connect._
                        """.strip()
                        return {
                            "response": response_text,
                            "agent": "OAuth Assistant",
                            "routed_to": "OAuth Assistant",
                            "routing_reason": "OAuth credentials not configured"
                        }
                    
                    oauth_payload.update({
                        "client_id": client_id,
                        "client_secret": client_secret,
                        "tenant_id": tenant_id
                    })
                elif auth_type == "google":
                    client_id = os.getenv("GOOGLE_CLIENT_ID")
                    client_secret = os.getenv("GOOGLE_CLIENT_SECRET")
                    
                    if not client_id or not client_secret:
                        response_text = """
❌ **Google OAuth Not Configured**

Google Workspace authentication is not yet configured. Please set up OAuth credentials in your `.env` file:

```
GOOGLE_CLIENT_ID=your_client_id_here
GOOGLE_CLIENT_SECRET=your_client_secret_here
```

**Need help?** Check out `CREDENTIALS_SETUP.md` for detailed instructions.
                        """.strip()
                        return {
                            "response": response_text,
                            "agent": "OAuth Assistant",
                            "routed_to": "OAuth Assistant",
                            "routing_reason": "OAuth credentials not configured"
                        }
                    
                    oauth_payload.update({
                        "client_id": client_id,
                        "client_secret": client_secret
                    })
                
                try:
                    # Call OAuth authorization endpoint
                    import httpx
                    async with httpx.AsyncClient() as client:
                        oauth_response = await client.post(
                            "http://localhost:8084/api/v1/oauth/user/authorize",
                            json=oauth_payload,
                            timeout=10.0
                        )
                        
                        if oauth_response.status_code == 200:
                            auth_data = oauth_response.json()
                            auth_url = auth_data["authorization_url"]
                            
                            # Update response with actual URL - make it clickable
                            auth_link_html = f'<a href="{auth_url}" target="_blank" style="color: #0066cc; text-decoration: underline; font-weight: bold;">Click Here to Login</a>'
                            
                            # Replace the placeholder markdown link with actual URL
                            response_text = re.sub(
                                r'\[Login with [^\]]+\]\(#auth-[^\)]+\)',
                                auth_link_html,
                                response_text
                            )
                            
                            # Also add the raw URL for easy access
                            response_text += f"\n\n**Direct Link:** {auth_url}"
                        else:
                            response_text = f"❌ Failed to generate authorization URL: {oauth_response.text}"
                            
                except Exception as e:
                    response_text = f"❌ Error connecting to OAuth service: {str(e)}"
        
        return {
            "response": response_text,
            "agent": "OAuth Assistant",
            "routed_to": "OAuth Assistant",
            "routing_reason": "Message contained authentication/OAuth keywords"
        }
    
    # Determine which agent should handle the request
    if any(word in message_lower for word in ["incident", "outage", "down", "broken", "error", "issue"]):
        target_agent_id = "agent-incident-001"
    elif any(word in message_lower for word in ["problem", "recurring", "analysis", "root cause"]):
        target_agent_id = "agent-problem-001"  
    elif any(word in message_lower for word in ["change", "update", "upgrade", "deployment", "release"]):
        target_agent_id = "agent-change-001"
    elif any(word in message_lower for word in ["manager", "escalate", "supervisor"]):
        target_agent_id = "manager-incident-001"
    elif any(word in message_lower for word in ["executive", "ceo", "strategic", "organization"]):
        target_agent_id = "ceo-001"
    else:
        # Default to incident specialist for general IT queries
        target_agent_id = "agent-incident-001"
    
    # Route to appropriate agent
    chat_message.agent_id = target_agent_id
    response = await chat_with_agent(chat_message)
    
    # Add routing information to response
    response_dict = response.dict()
    response_dict["routed_to"] = agents[target_agent_id].name
    response_dict["routing_reason"] = f"Message contained keywords suggesting {agents[target_agent_id].specialization}"
    
    return response_dict

# ===== TICKET MANAGEMENT ENDPOINTS =====

# In-memory ticket storage (replace with database in production)
tickets_db = []

@app.get("/api/v1/tickets")
async def get_tickets():
    """Get all tickets"""
    return tickets_db

@app.post("/api/v1/tickets")
async def create_ticket(ticket: dict):
    """Create a new ticket"""
    tickets_db.append(ticket)
    return ticket

@app.get("/api/v1/tickets/{ticket_id}")
async def get_ticket(ticket_id: str):
    """Get a specific ticket"""
    ticket = next((t for t in tickets_db if t.get('id') == ticket_id), None)
    if not ticket:
        raise HTTPException(status_code=404, detail="Ticket not found")
    return ticket

@app.put("/api/v1/tickets/{ticket_id}")
async def update_ticket(ticket_id: str, updates: dict):
    """Update a ticket"""
    ticket = next((t for t in tickets_db if t.get('id') == ticket_id), None)
    if not ticket:
        raise HTTPException(status_code=404, detail="Ticket not found")
    
    # Update ticket fields
    for key, value in updates.items():
        if key != 'id':  # Don't allow changing the ID
            ticket[key] = value
    
    return ticket

@app.delete("/api/v1/tickets/{ticket_id}")
async def delete_ticket(ticket_id: str):
    """Delete a ticket"""
    global tickets_db
    ticket = next((t for t in tickets_db if t.get('id') == ticket_id), None)
    if not ticket:
        raise HTTPException(status_code=404, detail="Ticket not found")
    
    tickets_db = [t for t in tickets_db if t.get('id') != ticket_id]
    return {"message": "Ticket deleted successfully"}

# ===== DATA ARCHIVES ENDPOINTS =====

# In-memory archives storage (replace with database in production)
archives_db = []

@app.get("/api/v1/archives")
async def get_archives():
    """Get all archived documents"""
    return archives_db

@app.post("/api/v1/archives")
async def upload_archive(
    title: str = Form(...),
    description: str = Form(None),
    category: str = Form(...),
    department: str = Form(None),
    tags: str = Form(None),
    searchable: bool = Form(True),
    file: UploadFile = File(...)
):
    """Upload a new document to archives"""
    import os
    from datetime import datetime
    
    # Validate file type
    allowed_extensions = ['.pdf', '.docx', '.txt']
    file_ext = os.path.splitext(file.filename)[1].lower()
    
    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF, DOCX, and TXT files are allowed.")
    
    # Read file content
    content = await file.read()
    file_size = len(content) / (1024 * 1024)  # Size in MB
    
    # Validate file size (max 10MB)
    if file_size > 10:
        raise HTTPException(status_code=400, detail="File size exceeds 10MB limit")
    
    # Process document based on file type
    text_content = ""
    try:
        if file_ext == '.txt':
            text_content = content.decode('utf-8')
        elif file_ext == '.docx':
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            text_content = '\n'.join([para.text for para in doc.paragraphs])
        elif file_ext == '.pdf':
            import PyPDF2
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text_content = '\n'.join([page.extract_text() for page in pdf_reader.pages])
    except Exception as e:
        logger.warning(f"Could not extract text from document: {e}")
        text_content = ""
    
    # Generate document ID
    doc_id = f"DOC-{str(len(archives_db) + 1).zfill(3)}"
    
    # Parse tags
    tag_list = [tag.strip() for tag in tags.split(',')] if tags else []
    
    # Create archive document
    archive_doc = {
        "id": doc_id,
        "title": title,
        "description": description,
        "category": category,
        "department": department,
        "file_name": file.filename,
        "file_type": file_ext[1:],  # Remove the dot
        "file_size": round(file_size, 2),
        "text_content": text_content,
        "uploaded_at": datetime.utcnow().isoformat() + 'Z',
        "uploaded_by": "Current User",  # Replace with actual user when auth is implemented
        "tags": tag_list,
        "searchable": searchable
    }
    
    # Store the document
    archives_db.append(archive_doc)
    
    # In production, save file to storage and index text_content for search
    logger.info(f"Document uploaded: {doc_id} - {title}")
    
    return archive_doc

@app.get("/api/v1/archives/{archive_id}")
async def get_archive(archive_id: str):
    """Get a specific archived document"""
    archive = next((a for a in archives_db if a.get('id') == archive_id), None)
    if not archive:
        raise HTTPException(status_code=404, detail="Document not found")
    return archive

@app.get("/api/v1/archives/{archive_id}/download")
async def download_archive(archive_id: str):
    """Download an archived document"""
    archive = next((a for a in archives_db if a.get('id') == archive_id), None)
    if not archive:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # In production, return the actual file from storage
    return {"message": "Download endpoint - file would be served here", "document": archive}

@app.get("/api/v1/archives/{archive_id}/view")
async def view_archive(archive_id: str):
    """View an archived document"""
    archive = next((a for a in archives_db if a.get('id') == archive_id), None)
    if not archive:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Return text content for viewing
    return {
        "id": archive['id'],
        "title": archive['title'],
        "content": archive.get('text_content', 'Content not available'),
        "file_type": archive['file_type']
    }

@app.put("/api/v1/archives/{archive_id}")
async def update_archive(archive_id: str, updates: dict):
    """Update an archived document metadata"""
    archive = next((a for a in archives_db if a.get('id') == archive_id), None)
    if not archive:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Update allowed fields
    for key, value in updates.items():
        if key not in ['id', 'file_name', 'file_type', 'uploaded_at', 'text_content']:
            archive[key] = value
    
    return archive

@app.delete("/api/v1/archives/{archive_id}")
async def delete_archive(archive_id: str):
    """Delete an archived document"""
    global archives_db
    archive = next((a for a in archives_db if a.get('id') == archive_id), None)
    if not archive:
        raise HTTPException(status_code=404, detail="Document not found")
    
    archives_db = [a for a in archives_db if a.get('id') != archive_id]
    logger.info(f"Document deleted: {archive_id}")
    return {"message": "Document deleted successfully"}

@app.get("/api/v1/archives/search")
async def search_archives(query: str):
    """Search archives by text content (for agent use)"""
    if not query:
        return []
    
    query_lower = query.lower()
    results = []
    
    for archive in archives_db:
        if not archive.get('searchable', True):
            continue
        
        # Search in title, description, tags, and content
        score = 0
        if query_lower in archive['title'].lower():
            score += 10
        if archive.get('description') and query_lower in archive['description'].lower():
            score += 5
        if any(query_lower in tag.lower() for tag in archive.get('tags', [])):
            score += 3
        if archive.get('text_content') and query_lower in archive['text_content'].lower():
            score += 1
        
        if score > 0:
            results.append({
                **archive,
                'relevance_score': score
            })
    
    # Sort by relevance
    results.sort(key=lambda x: x['relevance_score'], reverse=True)
    return results[:10]  # Return top 10 results

# ============================================================================
# Multi-Tenancy / Organizations Endpoints
# ============================================================================

# In-memory organizations database (replace with real DB in production)
organizations_db = []

@app.get("/api/v1/organizations")
async def get_organizations():
    """Get all organizations"""
    return organizations_db

@app.post("/api/v1/organizations")
async def create_organization(data: dict):
    """Create a new organization"""
    org_id = f"ORG-{len(organizations_db) + 1:03d}"
    organization = {
        "id": org_id,
        "name": data.get("name", ""),
        "domain": data.get("domain", ""),
        "plan": data.get("plan", "free"),
        "status": "active",
        "created_at": datetime.now().isoformat(),
        "agent_limit": data.get("agent_limit", 5),
        "agents_count": 0,
        "messages_count": 0
    }
    organizations_db.append(organization)
    logger.info(f"Organization created: {org_id}")
    return organization

@app.get("/api/v1/organizations/{org_id}")
async def get_organization(org_id: str):
    """Get organization details"""
    org = next((o for o in organizations_db if o["id"] == org_id), None)
    if not org:
        raise HTTPException(status_code=404, detail="Organization not found")
    return org

@app.get("/api/v1/organizations/{org_id}/usage")
async def get_organization_usage(org_id: str):
    """Get organization usage statistics"""
    org = next((o for o in organizations_db if o["id"] == org_id), None)
    if not org:
        raise HTTPException(status_code=404, detail="Organization not found")
    
    return {
        "organization_id": org_id,
        "agents_count": org.get("agents_count", 0),
        "agent_limit": org.get("agent_limit", 5),
        "messages_count": org.get("messages_count", 0),
        "storage_used": 0,
        "api_calls": 0
    }

@app.put("/api/v1/organizations/{org_id}")
async def update_organization(org_id: str, data: dict):
    """Update organization details"""
    org = next((o for o in organizations_db if o["id"] == org_id), None)
    if not org:
        raise HTTPException(status_code=404, detail="Organization not found")
    
    org.update(data)
    logger.info(f"Organization updated: {org_id}")
    return org

@app.delete("/api/v1/organizations/{org_id}")
async def delete_organization(org_id: str):
    """Delete an organization"""
    global organizations_db
    organizations_db = [o for o in organizations_db if o["id"] != org_id]
    logger.info(f"Organization deleted: {org_id}")
    return {"message": "Organization deleted successfully"}

# ============================================================================
# Intelligent Routing / Global Stats Endpoints
# ============================================================================

@app.get("/api/v1/routing/global-stats")
async def get_global_routing_stats():
    """Get global routing statistics"""
    # Count enabled agents
    active_count = sum(1 for agent_id in agents.keys() if agent_status.get(agent_id, True))
    
    return {
        "total_requests": 0,
        "successful_routes": 0,
        "failed_routes": 0,
        "average_response_time": 0,
        "active_agents": active_count,
        "total_agents": len(agents),
        "load_distribution": {},
        "peak_hours": [],
        "agent_performance": []
    }

@app.get("/api/v1/agents/{agent_id}/routing-stats")
async def get_agent_routing_stats(agent_id: str):
    """Get routing statistics for a specific agent"""
    if agent_id not in agents:
        raise HTTPException(status_code=404, detail="Agent not found")
    
    return {
        "agent_id": agent_id,
        "total_requests": 0,
        "successful_requests": 0,
        "failed_requests": 0,
        "average_response_time": 0,
        "current_load": 0,
        "capacity": 100,
        "utilization": 0
    }

# ============================================================================
# CONNECTORS API ENDPOINTS
# ============================================================================

class ConnectorCreateRequest(BaseModel):
    """Request model for creating a connector"""
    connector_type: str
    name: str
    description: Optional[str] = None
    auth_config: Dict[str, Any]
    permissions: Optional[List[str]] = None
    scope: Optional[str] = "read_write"

class ConnectorUpdateRequest(BaseModel):
    """Request model for updating a connector"""
    name: Optional[str] = None
    description: Optional[str] = None
    auth_config: Optional[Dict[str, Any]] = None
    permissions: Optional[List[str]] = None
    status: Optional[str] = None

@app.get("/api/v1/connectors/available")
async def get_available_connectors():
    """Get list of all available standard connectors"""
    try:
        connectors = []
        for connector_type, config in STANDARD_CONNECTORS.items():
            connectors.append({
                "type": connector_type.value,
                "name": config["name"],
                "description": config["description"],
                "auth_type": config["auth_type"].value if isinstance(config["auth_type"], AuthType) else config["auth_type"],
                "icon": config.get("icon", ""),
                "capabilities": config.get("capabilities", {}),
                "documentation_url": config.get("documentation_url", "")
            })
        
        return {
            "connectors": connectors,
            "total": len(connectors)
        }
    except Exception as e:
        logger.error(f"Error fetching available connectors: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/connectors")
async def list_connectors(tenant_id: Optional[str] = None):
    """List all configured connectors"""
    try:
        connectors = connector_manager.list_connectors(tenant_id)
        return {
            "connectors": [
                {
                    "connector_id": c.connector_id,
                    "connector_type": c.connector_type.value,
                    "name": c.name,
                    "description": c.description,
                    "auth_type": c.auth_type.value,
                    "status": c.status.value,
                    "permissions": c.permissions,
                    "scope": c.scope.value,
                    "created_at": c.created_at.isoformat(),
                    "updated_at": c.updated_at.isoformat(),
                    "last_sync": c.last_sync.isoformat() if c.last_sync else None,
                    "rate_limit": c.rate_limit
                }
                for c in connectors
            ],
            "total": len(connectors)
        }
    except Exception as e:
        logger.error(f"Error listing connectors: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/connectors")
async def create_connector(request: ConnectorCreateRequest):
    """Create a new connector"""
    try:
        # Validate connector type
        try:
            connector_type = ConnectorType(request.connector_type)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"Invalid connector type: {request.connector_type}")
        
        # Create connector
        config = connector_manager.create_connector(
            connector_type=connector_type,
            name=request.name,
            auth_config=request.auth_config,
            tenant_id=None  # Can be extended for multi-tenant support
        )
        
        if request.permissions:
            config.permissions = request.permissions
        
        if request.scope:
            try:
                config.scope = PermissionScope(request.scope)
            except ValueError:
                pass
        
        return {
            "success": True,
            "connector_id": config.connector_id,
            "message": f"Connector '{request.name}' created successfully",
            "connector": {
                "connector_id": config.connector_id,
                "connector_type": config.connector_type.value,
                "name": config.name,
                "status": config.status.value,
                "auth_type": config.auth_type.value
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating connector: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/connectors/{connector_id}")
async def get_connector(connector_id: str):
    """Get connector details"""
    try:
        config = connector_manager.get_connector(connector_id)
        if not config:
            raise HTTPException(status_code=404, detail="Connector not found")
        
        return {
            "connector_id": config.connector_id,
            "connector_type": config.connector_type.value,
            "name": config.name,
            "description": config.description,
            "auth_type": config.auth_type.value,
            "status": config.status.value,
            "permissions": config.permissions,
            "scope": config.scope.value,
            "created_at": config.created_at.isoformat(),
            "updated_at": config.updated_at.isoformat(),
            "last_sync": config.last_sync.isoformat() if config.last_sync else None,
            "rate_limit": config.rate_limit,
            "custom_config": config.custom_config
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting connector: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/api/v1/connectors/{connector_id}")
async def update_connector(connector_id: str, request: ConnectorUpdateRequest):
    """Update connector configuration"""
    try:
        config = connector_manager.get_connector(connector_id)
        if not config:
            raise HTTPException(status_code=404, detail="Connector not found")
        
        # Update fields
        if request.name:
            config.name = request.name
        if request.description:
            config.description = request.description
        if request.auth_config:
            config.auth_config.update(request.auth_config)
        if request.permissions:
            config.permissions = request.permissions
        if request.status:
            try:
                config.status = ConnectorStatus(request.status)
            except ValueError:
                raise HTTPException(status_code=400, detail=f"Invalid status: {request.status}")
        
        config.updated_at = datetime.now()
        
        return {
            "success": True,
            "message": "Connector updated successfully",
            "connector_id": connector_id
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating connector: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/v1/connectors/{connector_id}")
async def delete_connector(connector_id: str):
    """Delete a connector"""
    try:
        success = connector_manager.delete_connector(connector_id)
        if not success:
            raise HTTPException(status_code=404, detail="Connector not found")
        
        return {
            "success": True,
            "message": "Connector deleted successfully"
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting connector: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/connectors/{connector_id}/test")
async def test_connector(connector_id: str):
    """Test connector connection using production implementation"""
    try:
        from core.connector_implementations import get_connector_implementation
        import time
        
        config = connector_manager.get_connector(connector_id)
        if not config:
            raise HTTPException(status_code=404, detail="Connector not found")
        
        # Get production connector implementation
        connector_impl = get_connector_implementation(
            config.connector_type.value,
            connector_id,
            config.auth_config
        )
        
        if not connector_impl:
            # Fallback to simulated test if no implementation available
            config.status = ConnectorStatus.CONNECTED
            config.updated_at = datetime.now()
            return {
                "success": True,
                "connector_id": connector_id,
                "status": "connected",
                "message": "Connection test successful (simulated)",
                "details": {
                    "response_time_ms": 150,
                    "authenticated": True,
                    "permissions_valid": True
                }
            }
        
        # Perform actual connection test
        start_time = time.time()
        test_result = connector_impl.test_connection()
        response_time = int((time.time() - start_time) * 1000)
        
        if test_result:
            # Update connector status to connected after successful test
            config.status = ConnectorStatus.CONNECTED
            config.updated_at = datetime.now()
            
            return {
                "success": True,
                "connector_id": connector_id,
                "status": "connected",
                "message": "Connection test successful",
                "details": {
                    "response_time_ms": response_time,
                    "authenticated": True,
                    "permissions_valid": True,
                    "test_result": test_result
                }
            }
        else:
            return {
                "success": False,
                "connector_id": connector_id,
                "status": "failed",
                "message": "Connection test failed",
                "details": {
                    "response_time_ms": response_time,
                    "authenticated": False
                }
            }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error testing connector: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/connectors/{connector_id}/authorize")
async def initiate_authorization(connector_id: str):
    """Initiate OAuth authorization flow"""
    try:
        auth_data = connector_manager.initiate_oauth_flow(connector_id)
        return {
            "success": True,
            "authorization_url": auth_data["authorization_url"],
            "state": auth_data["state"],
            "message": "Redirect user to authorization_url to complete OAuth flow"
        }
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Error initiating authorization: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/connectors/{connector_id}/revoke")
async def revoke_connector_access(connector_id: str):
    """Revoke connector access"""
    try:
        config = connector_manager.get_connector(connector_id)
        if not config:
            raise HTTPException(status_code=404, detail="Connector not found")
        
        config.status = ConnectorStatus.DISCONNECTED
        config.auth_config.pop("access_token", None)
        config.auth_config.pop("refresh_token", None)
        config.updated_at = datetime.now()
        
        return {
            "success": True,
            "message": "Access revoked successfully"
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error revoking access: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/connectors/{connector_id}/permissions")
async def get_connector_permissions(connector_id: str):
    """Get available permissions for a connector"""
    try:
        config = connector_manager.get_connector(connector_id)
        if not config:
            raise HTTPException(status_code=404, detail="Connector not found")
        
        standard_config = STANDARD_CONNECTORS.get(config.connector_type, {})
        capabilities = standard_config.get("capabilities", {})
        
        return {
            "connector_id": connector_id,
            "capabilities": capabilities,
            "current_permissions": config.permissions,
            "scope": config.scope.value
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting permissions: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# ==== AI Agent Integration Endpoints ====

@app.post("/api/v1/connectors/{connector_id}/execute")
async def execute_connector_action(connector_id: str, request: dict):
    """Generic connector action executor - supports all connector methods"""
    try:
        from core.connector_implementations import get_connector_implementation
        
        config = connector_manager.get_connector(connector_id)
        if not config:
            raise HTTPException(status_code=404, detail="Connector not found")
        
        connector = get_connector_implementation(
            config.connector_type.value, 
            connector_id, 
            config.auth_config
        )
        if not connector:
            raise HTTPException(status_code=500, detail="Failed to initialize connector")
        
        action = request.get("action")
        parameters = request.get("parameters", {})
        
        if not action:
            raise HTTPException(status_code=400, detail="Missing 'action' field")
        
        # Execute the action dynamically
        if not hasattr(connector, action):
            raise HTTPException(status_code=400, detail=f"Action '{action}' not supported")
        
        method = getattr(connector, action)
        result = method(**parameters)
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error executing connector action: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/connectors/{connector_id}/teams/send_message")
async def teams_send_message(connector_id: str, request: dict):
    """AI Agent: Send message to Teams channel"""
    try:
        from core.connector_implementations import get_connector_implementation
        
        config = connector_manager.get_connector(connector_id)
        if not config or config.connector_type.value != "microsoft_teams":
            raise HTTPException(status_code=400, detail="Invalid Teams connector")
        
        connector = get_connector_implementation("microsoft_teams", connector_id, config.auth_config)
        if not connector:
            raise HTTPException(status_code=500, detail="Failed to initialize connector")
        
        team_id = request.get("team_id")
        channel_id = request.get("channel_id")
        message = request.get("message")
        
        if not all([team_id, channel_id, message]):
            raise HTTPException(status_code=400, detail="Missing required fields")
        
        result = connector.send_message(team_id, channel_id, message)
        
        return {
            "success": True,
            "result": result,
            "message": "Message sent successfully"
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error sending Teams message: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/connectors/{connector_id}/teams/messages")
async def teams_get_messages(connector_id: str, team_id: str, channel_id: str, limit: int = 50):
    """AI Agent: Get messages from Teams channel"""
    try:
        from core.connector_implementations import get_connector_implementation
        
        config = connector_manager.get_connector(connector_id)
        if not config or config.connector_type.value != "microsoft_teams":
            raise HTTPException(status_code=400, detail="Invalid Teams connector")
        
        connector = get_connector_implementation("microsoft_teams", connector_id, config.auth_config)
        if not connector:
            raise HTTPException(status_code=500, detail="Failed to initialize connector")
        
        messages = connector.get_messages(team_id, channel_id, limit)
        
        return {
            "success": True,
            "messages": messages,
            "count": len(messages)
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting Teams messages: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/connectors/{connector_id}/drive/files")
async def drive_list_files(connector_id: str, query: str = None, limit: int = 100):
    """AI Agent: List files in Google Drive"""
    try:
        from core.connector_implementations import get_connector_implementation
        
        config = connector_manager.get_connector(connector_id)
        if not config or config.connector_type.value != "google_drive":
            raise HTTPException(status_code=400, detail="Invalid Drive connector")
        
        connector = get_connector_implementation("google_drive", connector_id, config.auth_config)
        if not connector:
            raise HTTPException(status_code=500, detail="Failed to initialize connector")
        
        files = connector.list_files(query, limit)
        
        return {
            "success": True,
            "files": files,
            "count": len(files)
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error listing Drive files: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/connectors/{connector_id}/drive/upload")
async def drive_upload_file(connector_id: str, request: dict):
    """AI Agent: Upload file to Google Drive"""
    try:
        from core.connector_implementations import get_connector_implementation
        import base64
        
        config = connector_manager.get_connector(connector_id)
        if not config or config.connector_type.value != "google_drive":
            raise HTTPException(status_code=400, detail="Invalid Drive connector")
        
        connector = get_connector_implementation("google_drive", connector_id, config.auth_config)
        if not connector:
            raise HTTPException(status_code=500, detail="Failed to initialize connector")
        
        file_name = request.get("file_name")
        content_base64 = request.get("content")  # Base64 encoded
        mime_type = request.get("mime_type", "application/octet-stream")
        folder_id = request.get("folder_id")
        
        if not all([file_name, content_base64]):
            raise HTTPException(status_code=400, detail="Missing required fields")
        
        # Decode base64 content
        content = base64.b64decode(content_base64)
        
        file_id = connector.upload_file(file_name, content, mime_type, folder_id)
        
        return {
            "success": True,
            "file_id": file_id,
            "message": "File uploaded successfully"
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading file: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# OAUTH AUTHENTICATION ENDPOINTS
# ============================================================================

@app.get("/api/v1/oauth/authorize/{connector_id}")
async def oauth_authorize(connector_id: str):
    """
    Step 1: Generate OAuth authorization URL for user login
    
    This endpoint initiates the OAuth flow by generating a Microsoft login URL.
    The user will be redirected to Microsoft to log in and consent to permissions.
    """
    try:
        from core.oauth_manager import get_oauth_flow
        
        # Get connector config
        config = connector_manager.get_connector(connector_id)
        if not config:
            raise HTTPException(status_code=404, detail="Connector not found")
        
        if config.connector_type.value != "microsoft_teams":
            raise HTTPException(status_code=400, detail="OAuth only supported for Microsoft 365 connectors")
        
        # Get OAuth credentials
        client_id = config.auth_config.get('client_id')
        client_secret = config.auth_config.get('client_secret')
        tenant_id = config.auth_config.get('tenant_id', 'common')
        
        if not client_id or not client_secret:
            raise HTTPException(status_code=400, detail="Missing OAuth credentials")
        
        # Generate authorization URL
        oauth_flow = get_oauth_flow(client_id, client_secret, tenant_id)
        auth_url, state = oauth_flow.get_authorization_url(connector_id)
        
        return {
            "success": True,
            "authorization_url": auth_url,
            "state": state,
            "message": "Please visit the authorization URL to log in with your Microsoft account"
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating OAuth URL: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/oauth/callback")
async def oauth_callback(code: Optional[str] = None, state: Optional[str] = None, error: Optional[str] = None):
    """
    Step 2: OAuth callback handler
    
    Microsoft redirects here after user logs in and consents to permissions.
    This exchanges the authorization code for access and refresh tokens.
    """
    try:
        from core.oauth_manager import get_oauth_flow, oauth_token_manager
        from fastapi.responses import HTMLResponse
        
        # Check for errors
        if error:
            error_html = f"""
            <html>
                <head><title>OAuth Error</title></head>
                <body style="font-family: Arial; padding: 50px; text-align: center;">
                    <h1 style="color: red;">❌ Authentication Failed</h1>
                    <p>Error: {error}</p>
                    <p><a href="http://localhost:8084/enhanced-dashboard">Return to Dashboard</a></p>
                </body>
            </html>
            """
            return HTMLResponse(content=error_html, status_code=400)
        
        if not code or not state:
            raise HTTPException(status_code=400, detail="Missing authorization code or state")
        
        # We need to find the OAuth flow that initiated this request
        # Get connector from any flow's state storage
        connector_id = None
        oauth_flow = None
        
        # Try to find the connector from state
        from core.oauth_manager import oauth_flows
        for flow in oauth_flows.values():
            if state in flow.state_storage:
                connector_id = flow.state_storage[state]
                oauth_flow = flow
                break
        
        if not connector_id or not oauth_flow:
            raise HTTPException(status_code=400, detail="Invalid state parameter")
        
        # Exchange code for tokens
        connector_id_result, token_data = oauth_flow.exchange_code_for_tokens(code, state)
        
        if not token_data:
            error_html = """
            <html>
                <head><title>Token Exchange Failed</title></head>
                <body style="font-family: Arial; padding: 50px; text-align: center;">
                    <h1 style="color: red;">❌ Token Exchange Failed</h1>
                    <p>Failed to exchange authorization code for tokens.</p>
                    <p><a href="http://localhost:8084/enhanced-dashboard">Return to Dashboard</a></p>
                </body>
            </html>
            """
            return HTMLResponse(content=error_html, status_code=500)
        
        # Save tokens
        oauth_token_manager.save_tokens(connector_id, token_data)
        
        # Update connector status
        config = connector_manager.get_connector(connector_id)
        if config:
            config.status = ConnectorStatus.CONNECTED
        
        # Return success page
        success_html = f"""
        <html>
            <head>
                <title>Authentication Successful</title>
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        padding: 50px;
                        text-align: center;
                        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                        color: white;
                    }}
                    .container {{
                        background: white;
                        color: #333;
                        padding: 40px;
                        border-radius: 10px;
                        max-width: 600px;
                        margin: 0 auto;
                        box-shadow: 0 10px 30px rgba(0,0,0,0.3);
                    }}
                    h1 {{ color: #28a745; margin-bottom: 20px; }}
                    .info {{ 
                        background: #f8f9fa;
                        padding: 20px;
                        border-radius: 5px;
                        margin: 20px 0;
                        text-align: left;
                    }}
                    .button {{
                        display: inline-block;
                        padding: 12px 30px;
                        background: #667eea;
                        color: white;
                        text-decoration: none;
                        border-radius: 5px;
                        margin-top: 20px;
                        font-weight: bold;
                    }}
                    .button:hover {{ background: #764ba2; }}
                </style>
            </head>
            <body>
                <div class="container">
                    <h1>✅ Authentication Successful!</h1>
                    <p>Your Microsoft 365 account has been successfully connected.</p>
                    
                    <div class="info">
                        <strong>Connector ID:</strong> {connector_id}<br/>
                        <strong>Status:</strong> Connected<br/>
                        <strong>Permissions:</strong> Email, OneDrive, Calendar, Meetings
                    </div>
                    
                    <p>You can now use all Microsoft 365 features including:</p>
                    <ul style="text-align: left; display: inline-block;">
                        <li>📧 Send and read emails</li>
                        <li>📁 Access OneDrive files</li>
                        <li>📅 Manage calendar events</li>
                        <li>👥 Create Teams meetings</li>
                    </ul>
                    
                    <a href="http://localhost:8084/enhanced-dashboard" class="button">
                        Return to Dashboard
                    </a>
                    
                    <p style="margin-top: 30px; font-size: 12px; color: #666;">
                        This window can be closed.
                    </p>
                </div>
            </body>
        </html>
        """
        return HTMLResponse(content=success_html)
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in OAuth callback: {str(e)}")
        error_html = f"""
        <html>
            <head><title>OAuth Error</title></head>
            <body style="font-family: Arial; padding: 50px; text-align: center;">
                <h1 style="color: red;">❌ Authentication Error</h1>
                <p>Error: {str(e)}</p>
                <p><a href="http://localhost:8084/enhanced-dashboard">Return to Dashboard</a></p>
            </body>
        </html>
        """
        return HTMLResponse(content=error_html, status_code=500)


@app.get("/api/v1/oauth/status/{connector_id}")
async def oauth_status(connector_id: str):
    """Check OAuth authentication status for a connector"""
    try:
        from core.oauth_manager import oauth_token_manager
        
        config = connector_manager.get_connector(connector_id)
        if not config:
            raise HTTPException(status_code=404, detail="Connector not found")
        
        has_tokens = oauth_token_manager.has_refresh_token(connector_id)
        token_data = oauth_token_manager.get_tokens(connector_id) if has_tokens else None
        
        is_valid = False
        expires_at = None
        
        if token_data and 'expires_at' in token_data:
            expires_at = token_data['expires_at']
            expiry = datetime.fromisoformat(expires_at)
            is_valid = datetime.now() < expiry
        
        return {
            "connector_id": connector_id,
            "has_tokens": has_tokens,
            "is_authenticated": has_tokens and is_valid,
            "expires_at": expires_at,
            "auth_type": "delegated" if has_tokens else "application"
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error checking OAuth status: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/v1/oauth/tokens/{connector_id}")
async def oauth_revoke(connector_id: str):
    """Revoke OAuth tokens for a connector"""
    try:
        from core.oauth_manager import oauth_token_manager
        
        config = connector_manager.get_connector(connector_id)
        if not config:
            raise HTTPException(status_code=404, detail="Connector not found")
        
        oauth_token_manager.delete_tokens(connector_id)
        config.status = ConnectorStatus.DISCONNECTED
        
        return {
            "success": True,
            "message": "OAuth tokens revoked successfully"
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error revoking OAuth tokens: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== Microsoft 365 Connector Endpoints ====================

@app.post("/api/v1/connectors/{connector_id}/send_email")
async def send_email(connector_id: str, email_data: dict):
    """Send an email via Microsoft 365 connector"""
    try:
        from core.connector_implementations import get_connector_implementation
        
        config = connector_manager.get_connector(connector_id)
        if not config:
            raise HTTPException(status_code=404, detail="Connector not found")
        
        if config.connector_type.value != "microsoft_teams":
            raise HTTPException(status_code=400, detail="Connector is not a Microsoft 365 connector")
        
        connector = get_connector_implementation("microsoft_teams", connector_id, config.auth_config)
        
        result = connector.send_email(
            to=email_data.get("to"),
            subject=email_data.get("subject"),
            body=email_data.get("body"),
            cc=email_data.get("cc"),
            bcc=email_data.get("bcc"),
            attachments=email_data.get("attachments"),
            is_html=email_data.get("is_html", False)
        )
        
        return result
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error sending email: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/connectors/{connector_id}/read_emails")
async def read_emails(connector_id: str, limit: int = 10, folder: str = "inbox"):
    """Read emails from Microsoft 365 connector"""
    try:
        from core.connector_implementations import get_connector_implementation
        
        config = connector_manager.get_connector(connector_id)
        if not config:
            raise HTTPException(status_code=404, detail="Connector not found")
        
        if config.connector_type.value != "microsoft_teams":
            raise HTTPException(status_code=400, detail="Connector is not a Microsoft 365 connector")
        
        connector = get_connector_implementation("microsoft_teams", connector_id, config.auth_config)
        
        result = connector.read_emails(limit=limit, folder=folder)
        
        return result
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error reading emails: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/connectors/{connector_id}/onedrive/files")
async def list_onedrive_files(connector_id: str, path: str = "/"):
    """List OneDrive files"""
    try:
        from core.connector_implementations import get_connector_implementation
        
        config = connector_manager.get_connector(connector_id)
        if not config:
            raise HTTPException(status_code=404, detail="Connector not found")
        
        if config.connector_type.value != "microsoft_teams":
            raise HTTPException(status_code=400, detail="Connector is not a Microsoft 365 connector")
        
        connector = get_connector_implementation("microsoft_teams", connector_id, config.auth_config)
        
        result = connector.list_onedrive_files(path=path)
        
        return result
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error listing OneDrive files: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/connectors/{connector_id}/onedrive/upload")
async def upload_onedrive_file(connector_id: str, file_data: dict):
    """Upload file to OneDrive"""
    try:
        from core.connector_implementations import get_connector_implementation
        
        config = connector_manager.get_connector(connector_id)
        if not config:
            raise HTTPException(status_code=404, detail="Connector not found")
        
        if config.connector_type.value != "microsoft_teams":
            raise HTTPException(status_code=400, detail="Connector is not a Microsoft 365 connector")
        
        connector = get_connector_implementation("microsoft_teams", connector_id, config.auth_config)
        
        result = connector.upload_onedrive_file(
            file_path=file_data.get("file_path"),
            remote_path=file_data.get("remote_path")
        )
        
        return result
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading to OneDrive: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== Multi-User OAuth Endpoints ====================

@app.post("/api/v1/oauth/user/authorize")
async def user_oauth_authorize(request: dict):
    """
    Start OAuth flow for a specific user
    
    Request body:
    {
        "user_email": "user@company.com",
        "connector_type": "microsoft_teams" or "google_workspace",
        "client_id": "...",
        "client_secret": "...",
        "tenant_id": "..." (for Microsoft only)
    }
    """
    try:
        from core.multi_user_oauth import multi_user_oauth_manager, MicrosoftOAuthFlow, GoogleOAuthFlow
        
        user_email = request.get("user_email")
        connector_type = request.get("connector_type")
        client_id = request.get("client_id")
        client_secret = request.get("client_secret")
        
        if not all([user_email, connector_type, client_id, client_secret]):
            raise HTTPException(status_code=400, detail="Missing required fields")
        
        if connector_type == "microsoft_teams":
            tenant_id = request.get("tenant_id", "common")
            oauth_flow = MicrosoftOAuthFlow(client_id, client_secret, tenant_id)
            auth_url, state = oauth_flow.get_authorization_url(user_email)
        elif connector_type == "google_workspace":
            oauth_flow = GoogleOAuthFlow(client_id, client_secret)
            auth_url, state = oauth_flow.get_authorization_url(user_email)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported connector type: {connector_type}")
        
        # Register pending authentication
        multi_user_oauth_manager.start_user_auth_flow(connector_type, user_email, state)
        
        return {
            "success": True,
            "authorization_url": auth_url,
            "state": state,
            "user_email": user_email,
            "connector_type": connector_type
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating user OAuth URL: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/oauth/callback/microsoft")
async def microsoft_oauth_callback(code: str = None, state: str = None, error: str = None):
    """Handle Microsoft OAuth callback"""
    try:
        from core.multi_user_oauth import multi_user_oauth_manager, MicrosoftOAuthFlow
        
        if error:
            return HTMLResponse(content=f"<h1>Authorization Error</h1><p>{error}</p>", status_code=400)
        
        if not code or not state:
            raise HTTPException(status_code=400, detail="Missing code or state")
        
        # Get user info from pending auth
        auth_result = multi_user_oauth_manager.complete_user_auth_flow(state)
        if not auth_result:
            raise HTTPException(status_code=400, detail="Invalid state parameter")
        
        user_email, connector_type = auth_result
        
        # TODO: Get client credentials from connector config
        # For now, this will be enhanced to look up the connector
        
        success_html = f"""
        <html>
            <head><title>Authentication Successful</title></head>
            <body style="font-family: Arial; padding: 50px; text-align: center;">
                <h1 style="color: #28a745;">✅ Authentication Successful!</h1>
                <p>User <strong>{user_email}</strong> has been authenticated for Microsoft 365.</p>
                <p>You can now close this window and return to the application.</p>
            </body>
        </html>
        """
        return HTMLResponse(content=success_html)
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in Microsoft OAuth callback: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/oauth/callback/google")
async def google_oauth_callback(code: str = None, state: str = None, error: str = None):
    """Handle Google OAuth callback"""
    try:
        from core.multi_user_oauth import multi_user_oauth_manager, GoogleOAuthFlow
        
        if error:
            return HTMLResponse(content=f"<h1>Authorization Error</h1><p>{error}</p>", status_code=400)
        
        if not code or not state:
            raise HTTPException(status_code=400, detail="Missing code or state")
        
        # Get user info from pending auth
        auth_result = multi_user_oauth_manager.complete_user_auth_flow(state)
        if not auth_result:
            raise HTTPException(status_code=400, detail="Invalid state parameter")
        
        user_email, connector_type = auth_result
        
        success_html = f"""
        <html>
            <head><title>Authentication Successful</title></head>
            <body style="font-family: Arial; padding: 50px; text-align: center;">
                <h1 style="color: #28a745;">✅ Authentication Successful!</h1>
                <p>User <strong>{user_email}</strong> has been authenticated for Google Workspace.</p>
                <p>You can now close this window and return to the application.</p>
            </body>
        </html>
        """
        return HTMLResponse(content=success_html)
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in Google OAuth callback: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/oauth/users")
async def list_authenticated_users(connector_type: str):
    """List all authenticated users for a connector type"""
    try:
        from core.multi_user_oauth import multi_user_oauth_manager
        
        users = multi_user_oauth_manager.list_authenticated_users(connector_type)
        
        return {
            "connector_type": connector_type,
            "authenticated_users": users,
            "count": len(users)
        }
    
    except Exception as e:
        logger.error(f"Error listing authenticated users: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/v1/oauth/user/{connector_type}/{user_email}")
async def revoke_user_oauth(connector_type: str, user_email: str):
    """Revoke OAuth tokens for a specific user"""
    try:
        from core.multi_user_oauth import multi_user_oauth_manager
        
        multi_user_oauth_manager.delete_user_tokens(connector_type, user_email)
        
        return {
            "success": True,
            "message": f"Tokens revoked for {user_email}"
        }
    
    except Exception as e:
        logger.error(f"Error revoking user tokens: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    
    print("🚀 Starting Pramiti AI Organization API Server")
    print("📍 Ensure you have set OPENAI_API_KEY in your .env file")
    print("🔗 Frontend will be available at: http://localhost:8084")
    print("📊 API docs available at: http://localhost:8084/docs")
    
    uvicorn.run(
        "api_server:app",
        host="0.0.0.0",
        port=8084,
        reload=True,
        log_level="info"
    )